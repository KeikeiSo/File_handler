"""
program discription: group of functions that extract text from pdf,
repeat the sentences, and produce a new pdf

note: change the x and y tolerance if extraction produces error

author: Qiqi Su
"""

""" import statement """
import pdfplumber
from pdfplumber.utils import to_list, cluster_objects, collate_line
from collections import defaultdict
import re
from gtts import gTTS
  
# This module is imported so that we can 
# play the converted audio
import os
from docx import Document
  

""" Gobal variables """
# you should change it if the program does not work properly

# x represent horizontal distance between characters
# y represent vertical distance between characters
# adjust the number for better text extraction from the pdf
DEFAULT_X_TOLERANCE = 2
DEFAULT_Y_TOLERANCE = 3


""" functions """

# return a list of sentences from the pdf
def extract_sentences(
    page,
    x_tolerance=DEFAULT_X_TOLERANCE,
    y_tolerance=DEFAULT_Y_TOLERANCE,
):
    # preprocess the list of char objects
    chars = page.chars
    chars = to_list(chars)

    if len(chars) == 0:
        return None

    # remove the titles
    # (by only accept most common size chars)
    groups = defaultdict(list)
    # only get characters above the line if there is one
    delimiter = 0
    for line in page.lines:
        if line["linewidth"] >= 1:
            delimiter = line["top"]
            
    if delimiter != 0:
        for char in chars:
            if char["top"] > delimiter:
                groups[char["size"]].append(char)
    else:
        for char in chars:
            groups[char["size"]].append(char)

    lists = groups.values()
    content = []
    for l in lists:
        if len(l) > len(content):
            content = l

    # grouping different lines
    doctop_clusters = cluster_objects(content, "doctop", y_tolerance)

    # get each line by joinning the chars into words
    lines = (collate_line(line_chars, x_tolerance)
             for line_chars in doctop_clusters)

    coll = "\n".join(lines)
    sentences = re.split('；|。', coll)

    return sentences

def extract_content(
    page,
    x_tolerance=DEFAULT_X_TOLERANCE,
    y_tolerance=DEFAULT_Y_TOLERANCE,
):
    # preprocess the list of char objects
    chars = page.chars
    chars = to_list(chars)

    if len(chars) == 0:
        return None

    # only get characters above the line if there is one
    delimiter = 0
    for line in page.lines:
        if line["linewidth"] >= 1:
            delimiter = line["top"]
    
    content = []
    """ changed for chinese pdf, delete afterward """
    if delimiter != 0:
        for char in chars:
            #if char["bottom"] < delimiter:
            if char["top"] > delimiter:
                content.append(char)
    else:
        content = chars

    # grouping different lines
    doctop_clusters = cluster_objects(content, "doctop", y_tolerance)

    # get each line by joinning the chars into words
    lines = (collate_line(line_chars, x_tolerance)
             for line_chars in doctop_clusters)

    coll = "\n".join(lines)
    
    content = re.split('；|。', coll)

    return content

# helper function for sentence_repeater
# return the starting index (-1 if not found)
def search_sentence(str1, str2):
    try:
        x = re.search(str1, str2)
        if x:
            return x.start()
        else:
            s = str1.split('\n')
            y = re.search(s[0], str2)
            if y:
                return y.start()
            else:
                return -1
    except Exception:
        s = str1.split('\n')
        try:
            y = re.search(s[0], str2)
            if y:
                return y.start()
            else:
                return -1
        except Exception:
            return -1

# return a list of tokenized text with each sentence repeated
def sentence_repeater(page, frequency):
    context = extract_content(page)
    sentences = extract_sentences(page)
    

    document1 = Document()
    document1.add_paragraph("\n".join(context))
    document1.save("C:/Users/calli/Testing/text.docx")

    document2 = Document()
    document2.add_paragraph("\n".join(sentences))
    document2.save("C:/Users/calli/Testing/sentences.docx")


    index = 0

    # find the place first sentence match
    indicator = sentences[0]
    while (index < len(context) - 1):
        if indicator == context[index]:
            break
        else:
            if search_sentence(indicator, context[index]) == -1:
                index += 1
            else: 
                break
    # repeat each sentence and add them back
    for sentence in sentences:
        if sentence == context[index]:
            to_insert = sentence + "；\n"
            to_insert = frequency * to_insert
            context[index] = to_insert
            index += 1
        else:
            start = search_sentence(sentence, context[index])
            if start != -1:
                to_insert = sentence + "；\n"
                to_insert = frequency * to_insert
                context[index] = context[index][:start] + to_insert \
                    + context[index][start + len(sentence):]
                index += 1
            else:
                break

    return context


def split_for(str, pos):
    t = len(str) // pos
    if t == 0:
        return str
    for i in range(t):
        d = pos*(i+1)-1
        while(str[d] != " "):
            d += 1
            if d >= len(str)-1:
                break
        str = str[:d] + "\n" + str[d:]
    return str

def trim_text(text, pos):
    new_lines = []
    lines = text.split('\n')
    for line in lines:
        new_lines.append(split_for(line, pos))
    return "\n".join(new_lines)


# return a list of content of pages in the pdf with sentences repeated
def pdf_to_text(pages, start, end, frequency):
    sections = []
    pages = pages[start:end]
    for page in pages:
        section = sentence_repeater(page, frequency)
        sections.append(section)
    
    """ print to text file, delete afterward
    content = []
    for section in sections:
        content.append('\n'.join(section))
    text = '\n'.join(content) 
    text = trim_text(text, 90)

    text_file = open("C:/Users/calli/Testing/chinesetext.txt", "w")
    text_file.write(text)
    text_file.close()
    """

    return sections

def get_pages(filepath):
    pdf = pdfplumber.open(filepath)
    return pdf.pages

def pdf_to_docx(pages, start, end, frequency, outfile):
    sections = pdf_to_text(pages, start, end, frequency)
    document = Document()
    for section in sections:
        document.add_paragraph("\n".join(section))
    document.save(outfile)

    
""" main """
if __name__ == '__main__':
    pages = get_pages("C:/Users/calli/Testing/Cap 622 PDF (01-08-2019) (Traditional Chinese).pdf")
    sections = pdf_to_text(pages, 85, 86, 2)
    """
    content = []
    for section in sections:
        content.append('\n'.join(section))
    mytext = '\n'.join(content)
    """
    document = Document()
    for section in sections:
        document.add_paragraph("\n".join(section))
    document.save("C:/Users/calli/Testing/repeated.docx")


